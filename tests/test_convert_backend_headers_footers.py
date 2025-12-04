"""Tests for header and footer extraction from DOCX/ODT documents.

Headers and footers are page-level elements in Word/ODT documents that
pandoc does NOT extract (GitHub issue #5211). This test module:

1. Documents the current pandoc limitation
2. Tests python-docx based extraction as an alternative
3. Validates the integration into the conversion pipeline

Note: Extracting headers/footers is a HIGH PRIORITY feature goal.
"""
from __future__ import annotations

import io
from pathlib import Path
from zipfile import ZipFile

import pytest

try:
    from docx import Document as DocxDocument
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False

from convert_backend import convert_service as conv_service
from convert_backend.convert_service import convert_one
from convert_backend.convert_types import ConversionOptions


FIXTURE_DIR = Path(__file__).resolve().parents[1] / "tests" / "fixtures" / "converter"
HEADER_FOOTER_FIXTURE = FIXTURE_DIR / "header_footer_sample.docx"

# Expected content in our test fixture - use unique markers not found in body
EXPECTED_HEADER = "[DOCX-HDR-001] TinyUtils Official Header"
EXPECTED_FOOTER = "[DOCX-FTR-001] Confidential Document Footer"
# Unique substring to check for in tests
HEADER_MARKER = "DOCX-HDR-001"
FOOTER_MARKER = "DOCX-FTR-001"


def extract_headers_footers_docx(docx_path: Path) -> dict:
    """Extract headers and footers from a DOCX file using python-docx.

    Returns a dict with:
        headers: list of header texts (one per section)
        footers: list of footer texts (one per section)
    """
    if not HAS_PYTHON_DOCX:
        return {"headers": [], "footers": [], "error": "python-docx not installed"}

    doc = DocxDocument(docx_path)
    headers = []
    footers = []

    for section in doc.sections:
        # Extract header text
        if section.header and section.header.paragraphs:
            h_text = "\n".join(p.text for p in section.header.paragraphs if p.text.strip())
            if h_text:
                headers.append(h_text)

        # Extract footer text
        if section.footer and section.footer.paragraphs:
            f_text = "\n".join(p.text for p in section.footer.paragraphs if p.text.strip())
            if f_text:
                footers.append(f_text)

    return {"headers": headers, "footers": footers}


@pytest.mark.skipif(not HEADER_FOOTER_FIXTURE.exists(), reason="Header/footer fixture missing")
def test_docx_fixture_has_header_and_footer() -> None:
    """Verify our test fixture actually has headers and footers."""
    if not HAS_PYTHON_DOCX:
        pytest.skip("python-docx not installed")

    result = extract_headers_footers_docx(HEADER_FOOTER_FIXTURE)

    assert len(result["headers"]) >= 1, "fixture should have at least one header"
    assert len(result["footers"]) >= 1, "fixture should have at least one footer"
    assert EXPECTED_HEADER in result["headers"][0], f"expected header text not found"
    assert EXPECTED_FOOTER in result["footers"][0], f"expected footer text not found"


@pytest.mark.skipif(not HEADER_FOOTER_FIXTURE.exists(), reason="Header/footer fixture missing")
def test_pandoc_does_not_extract_headers_footers() -> None:
    """Document that pandoc (by default) ignores headers and footers.

    This test exists to document the pandoc limitation (GitHub #5211).
    Headers and footers are page-level elements that pandoc does not
    extract during DOCX→Markdown conversion.
    """
    raw = HEADER_FOOTER_FIXTURE.read_bytes()
    opts = ConversionOptions()
    result = convert_one(
        input_bytes=raw,
        name=HEADER_FOOTER_FIXTURE.name,
        targets=["md"],
        from_format="docx",
        options=opts,
    )

    assert result.error is None, f"unexpected conversion error: {result.error}"

    md_art = next((a for a in result.outputs or [] if a.target == "md"), None)
    assert md_art is not None, "no Markdown artifact produced"

    md_text = md_art.data.decode("utf-8")

    # Document the limitation: pandoc does NOT include headers/footers
    # This assertion documents current behavior - if it starts failing,
    # pandoc has added header/footer support and we can update our pipeline!
    assert HEADER_MARKER not in md_text, (
        "pandoc unexpectedly extracted header text - check if pandoc added header support!"
    )
    assert FOOTER_MARKER not in md_text, (
        "pandoc unexpectedly extracted footer text - check if pandoc added footer support!"
    )


@pytest.mark.skipif(not HEADER_FOOTER_FIXTURE.exists(), reason="Header/footer fixture missing")
def test_python_docx_can_extract_headers_footers() -> None:
    """Verify python-docx can extract headers/footers that pandoc misses."""
    if not HAS_PYTHON_DOCX:
        pytest.skip("python-docx not installed")

    result = extract_headers_footers_docx(HEADER_FOOTER_FIXTURE)

    # python-docx SHOULD be able to extract these
    assert len(result["headers"]) >= 1, "python-docx should extract headers"
    assert len(result["footers"]) >= 1, "python-docx should extract footers"

    # Verify actual content
    all_headers = " ".join(result["headers"])
    all_footers = " ".join(result["footers"])

    assert HEADER_MARKER in all_headers, "header content should be extracted"
    assert FOOTER_MARKER in all_footers, "footer content should be extracted"


@pytest.mark.skipif(not HEADER_FOOTER_FIXTURE.exists(), reason="Header/footer fixture missing")
def test_combined_output_with_headers_footers() -> None:
    """Test that we can combine pandoc output with extracted headers/footers.

    This demonstrates the approach for augmenting pandoc's output with
    header/footer content extracted via python-docx.
    """
    if not HAS_PYTHON_DOCX:
        pytest.skip("python-docx not installed")

    # Step 1: Get pandoc's markdown output
    raw = HEADER_FOOTER_FIXTURE.read_bytes()
    opts = ConversionOptions()
    result = convert_one(
        input_bytes=raw,
        name=HEADER_FOOTER_FIXTURE.name,
        targets=["md"],
        from_format="docx",
        options=opts,
    )

    assert result.error is None
    md_art = next((a for a in result.outputs or [] if a.target == "md"), None)
    pandoc_md = md_art.data.decode("utf-8")

    # Step 2: Extract headers/footers with python-docx
    hf = extract_headers_footers_docx(HEADER_FOOTER_FIXTURE)

    # Step 3: Combine into augmented output
    header_section = ""
    if hf["headers"]:
        header_section = "> **Document Header:** " + " | ".join(hf["headers"]) + "\n\n"

    footer_section = ""
    if hf["footers"]:
        footer_section = "\n\n---\n\n> **Document Footer:** " + " | ".join(hf["footers"])

    combined_md = header_section + pandoc_md + footer_section

    # Verify the combined output has both body AND header/footer content
    assert "Section 1" in combined_md, "body content should be preserved"
    assert HEADER_MARKER in combined_md, "header should now be in output"
    assert FOOTER_MARKER in combined_md, "footer should now be in output"


# Future: Integration test for when header/footer extraction is added to pipeline
@pytest.mark.skip(reason="Feature not yet implemented in pipeline")
def test_conversion_options_include_headers_footers() -> None:
    """Test that ConversionOptions.include_headers_footers flag works.

    When implemented, this option should:
    1. Extract headers/footers using python-docx (for DOCX) or similar for ODT
    2. Prepend headers and append footers to the markdown output
    3. Format them distinctively (blockquotes, horizontal rules, etc.)
    """
    raw = HEADER_FOOTER_FIXTURE.read_bytes()
    opts = ConversionOptions()
    # Future: opts.include_headers_footers = True

    result = convert_one(
        input_bytes=raw,
        name=HEADER_FOOTER_FIXTURE.name,
        targets=["md"],
        from_format="docx",
        options=opts,
    )

    md_art = next((a for a in result.outputs or [] if a.target == "md"), None)
    md_text = md_art.data.decode("utf-8")

    # When implemented, these should pass:
    assert HEADER_MARKER in md_text, "header should be in output when flag is set"
    assert FOOTER_MARKER in md_text, "footer should be in output when flag is set"
