#!/usr/bin/env python3
"""Create DOCX fixture for text alignment and color testing.

Text alignment and colors have NO pandoc support - completely lost:
- Text color (RGB, theme colors) → Lost
- Highlight color → Lost
- Text alignment (left/center/right/justify) → Lost
- Only text content is preserved

⚠️ Will be implemented via Google Cloud with LibreOffice HTML export.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def main():
    """Create alignment_color_sample.docx fixture."""
    doc = Document()

    doc.add_heading('ALIGN-COLOR-TEST-001 Alignment & Color Test', level=1)

    doc.add_paragraph(
        'This document tests text alignment and color preservation for future Google Cloud implementation.'
    )

    # Text colors
    doc.add_heading('COLOR-TEST-001 Text Colors', level=2)

    p1 = doc.add_paragraph('COLOR-TEXT-001 This paragraph contains colored text: ')
    p1.add_run('red text').font.color.rgb = RGBColor(255, 0, 0)
    p1.add_run(', ')
    p1.add_run('blue text').font.color.rgb = RGBColor(0, 0, 255)
    p1.add_run(', ')
    p1.add_run('green text').font.color.rgb = RGBColor(0, 255, 0)
    p1.add_run('.')

    # Multiple colors in one paragraph
    doc.add_heading('COLOR-MULTI-001 Multiple Colors', level=2)
    p2 = doc.add_paragraph()
    p2.add_run('COLOR-RAINBOW-001 ').font.color.rgb = RGBColor(0, 0, 0)
    p2.add_run('Rainbow ').font.color.rgb = RGBColor(255, 0, 0)
    p2.add_run('colored ').font.color.rgb = RGBColor(255, 165, 0)
    p2.add_run('text ').font.color.rgb = RGBColor(255, 255, 0)
    p2.add_run('paragraph ').font.color.rgb = RGBColor(0, 255, 0)
    p2.add_run('with ').font.color.rgb = RGBColor(0, 0, 255)
    p2.add_run('many ').font.color.rgb = RGBColor(75, 0, 130)
    p2.add_run('colors.').font.color.rgb = RGBColor(148, 0, 211)

    # Text alignment - left
    doc.add_heading('ALIGN-LEFT-001 Left Alignment', level=2)
    p3 = doc.add_paragraph('ALIGN-LEFT-TEXT-001 This paragraph is left-aligned (default).')
    p3.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Text alignment - center
    doc.add_heading('ALIGN-CENTER-001 Center Alignment', level=2)
    p4 = doc.add_paragraph('ALIGN-CENTER-TEXT-001 This paragraph is center-aligned.')
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Text alignment - right
    doc.add_heading('ALIGN-RIGHT-001 Right Alignment', level=2)
    p5 = doc.add_paragraph('ALIGN-RIGHT-TEXT-001 This paragraph is right-aligned.')
    p5.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Text alignment - justify
    doc.add_heading('ALIGN-JUSTIFY-001 Justified Alignment', level=2)
    p6 = doc.add_paragraph(
        'ALIGN-JUSTIFY-TEXT-001 This paragraph is justified. '
        'Justified text spreads words across the full width of the page, '
        'creating straight edges on both left and right margins.'
    )
    p6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Combined: color + alignment
    doc.add_heading('COMBO-001 Combined Color & Alignment', level=2)
    p7 = doc.add_paragraph()
    p7.add_run('COMBO-TEXT-001 ').font.color.rgb = RGBColor(0, 0, 0)
    p7.add_run('This centered text is also colored red.').font.color.rgb = RGBColor(255, 0, 0)
    p7.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Expected behavior
    doc.add_heading('Expected Behavior', level=2)

    doc.add_paragraph(
        '⚠️ PANDOC LIMITATION: Text colors and alignment are COMPLETELY LOST. '
        'Only plain text content is preserved.'
    )

    doc.add_paragraph(
        '✅ FUTURE: Will be implemented via Google Cloud with LibreOffice HTML export '
        'to preserve CSS styling (color, text-align properties).'
    )

    doc.add_paragraph(
        'Tests validate: 1) Content is preserved, 2) Markers present, '
        '3) Infrastructure ready for GCloud LibreOffice implementation.'
    )

    output_path = Path(__file__).resolve().parents[1] / 'tests' / 'fixtures' / 'converter' / 'alignment_color_sample.docx'
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f'✅ Created: {output_path}')


if __name__ == '__main__':
    main()
