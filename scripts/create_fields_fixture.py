#!/usr/bin/env python3
"""Create DOCX fixture for fields and dynamic content testing.

Fields have LIMITED pandoc support - they are converted to static text only.
Field codes and update functionality are lost.

⚠️ Will be implemented via Google Cloud with LibreOffice/python-uno for field extraction.
"""
from __future__ import annotations

from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def main():
    """Create fields_sample.docx fixture."""
    doc = Document()

    doc.add_heading('FIELDS-TEST-001 Dynamic Fields Test', level=1)

    doc.add_paragraph(
        'This document tests dynamic field preservation for future Google Cloud implementation.'
    )

    # PAGE field
    doc.add_heading('FIELD-PAGE-001 Page Number Field', level=2)
    p1 = doc.add_paragraph('FIELD-PAGE-TEXT-001 This paragraph contains a page number field: ')
    # Note: python-docx doesn't support field insertion directly, so we add static text
    # The real fixture should be created manually in Word with actual fields
    p1.add_run('Page 1').bold = True
    p1.add_run(' (This would be a PAGE field in Word)')

    # DATE field
    doc.add_heading('FIELD-DATE-001 Date Field', level=2)
    p2 = doc.add_paragraph('FIELD-DATE-TEXT-001 Document date field: ')
    p2.add_run(datetime.now().strftime('%Y-%m-%d')).bold = True
    p2.add_run(' (This would be a DATE field in Word)')

    # TOC field
    doc.add_heading('FIELD-TOC-001 Table of Contents', level=2)
    doc.add_paragraph('FIELD-TOC-TEXT-001 Table of Contents:')
    doc.add_paragraph('1. Introduction ......... 1')
    doc.add_paragraph('2. Methods .............. 3')
    doc.add_paragraph('3. Results .............. 5')
    doc.add_paragraph('(This would be a TOC field in Word)')

    # FILENAME field
    doc.add_heading('FIELD-FILENAME-001 Filename Field', level=2)
    p3 = doc.add_paragraph('FIELD-FILENAME-TEXT-001 Current file: ')
    p3.add_run('fields_sample.docx').bold = True
    p3.add_run(' (This would be a FILENAME field in Word)')

    # AUTHOR field
    doc.add_heading('FIELD-AUTHOR-001 Author Field', level=2)
    p4 = doc.add_paragraph('FIELD-AUTHOR-TEXT-001 Author: ')
    p4.add_run('Test Author').bold = True
    p4.add_run(' (This would be an AUTHOR field in Word)')

    # HYPERLINK field
    doc.add_heading('FIELD-HYPERLINK-001 Hyperlink Field', level=2)
    p5 = doc.add_paragraph('FIELD-HYPERLINK-TEXT-001 Visit: ')
    p5.add_run('https://example.com').underline = True
    p5.add_run(' (This would be a HYPERLINK field in Word)')

    # SEQ field
    doc.add_heading('FIELD-SEQ-001 Sequence Field', level=2)
    doc.add_paragraph('FIELD-SEQ-TEXT-001 Figure 1: First figure')
    doc.add_paragraph('FIELD-SEQ-TEXT-002 Figure 2: Second figure')
    doc.add_paragraph('FIELD-SEQ-TEXT-003 Figure 3: Third figure')
    doc.add_paragraph('(These would be SEQ fields in Word)')

    # Expected behavior
    doc.add_heading('Expected Behavior', level=2)

    doc.add_paragraph(
        '⚠️ PANDOC LIMITATION: Fields are converted to static text. '
        'Field codes and update functionality are lost.'
    )

    doc.add_paragraph(
        '✅ FUTURE: Will be implemented via Google Cloud with LibreOffice/python-uno '
        'to extract and preserve field codes.'
    )

    doc.add_paragraph(
        'Tests validate: 1) Content is preserved (as static text), '
        '2) Field markers are present, 3) Infrastructure ready for GCloud field extraction.'
    )

    output_path = Path(__file__).resolve().parents[1] / 'tests' / 'fixtures' / 'converter' / 'fields_sample.docx'
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f'✅ Created: {output_path}')
    print('⚠️  NOTE: For real field testing, manually create DOCX with actual Word fields.')


if __name__ == '__main__':
    main()
