#!/usr/bin/env python3
"""Create a DOCX fixture to test typography preservation.

Typography features are converted by pandoc as:
- Heading levels → Markdown heading levels (# ## ###)
- Text colors → HTML spans with color styles
- Text alignment → HTML divs with align/style attributes
- Highlighting → HTML spans with background color
- Font family/size → NOT preserved (limitation)

This fixture creates a document with comprehensive typography features.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.shared import RGBColor


def main():
    """Create typography_sample.docx fixture."""
    doc = Document()

    # Title
    doc.add_heading('TYPO-TEST-001 Typography Test', level=1)

    doc.add_paragraph(
        'This document tests typography preservation through DOCX→Markdown conversions.'
    )

    # Section 1: Font Hierarchy (Heading Levels)
    doc.add_heading('TYPO-HEADINGS-001 Font Hierarchy', level=2)

    doc.add_paragraph(
        'Pandoc converts DOCX heading levels to Markdown heading syntax.'
    )

    doc.add_heading('TYPO-H1-001 Level 1 Heading', level=1)
    doc.add_heading('TYPO-H2-001 Level 2 Heading', level=2)
    doc.add_heading('TYPO-H3-001 Level 3 Heading', level=3)
    doc.add_heading('TYPO-H4-001 Level 4 Heading', level=4)
    doc.add_heading('TYPO-H5-001 Level 5 Heading', level=5)

    # Section 2: Text Colors
    doc.add_heading('TYPO-COLORS-001 Text Colors', level=2)

    doc.add_paragraph(
        'Text colors are converted to HTML spans with inline styles.'
    )

    p_colors = doc.add_paragraph('TYPO-COLOR-RED-001 ')
    run_red = p_colors.add_run('Red text')
    run_red.font.color.rgb = RGBColor(255, 0, 0)

    p_colors.add_run(', TYPO-COLOR-BLUE-001 ')
    run_blue = p_colors.add_run('Blue text')
    run_blue.font.color.rgb = RGBColor(0, 0, 255)

    p_colors.add_run(', TYPO-COLOR-GREEN-001 ')
    run_green = p_colors.add_run('Green text')
    run_green.font.color.rgb = RGBColor(0, 128, 0)

    # Section 3: Text Alignment
    doc.add_heading('TYPO-ALIGN-001 Text Alignment', level=2)

    doc.add_paragraph(
        'Text alignment is converted to HTML attributes or styles.'
    )

    p_left = doc.add_paragraph('TYPO-ALIGN-LEFT-001 Left aligned text')
    p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p_center = doc.add_paragraph('TYPO-ALIGN-CENTER-001 Center aligned text')
    p_center.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p_right = doc.add_paragraph('TYPO-ALIGN-RIGHT-001 Right aligned text')
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    p_justify = doc.add_paragraph('TYPO-ALIGN-JUSTIFY-001 Justified text with enough content to span multiple lines and demonstrate full justification.')
    p_justify.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Section 4: Highlighting
    doc.add_heading('TYPO-HIGHLIGHT-001 Text Highlighting', level=2)

    doc.add_paragraph(
        'Highlighting is converted to HTML spans with background colors.'
    )

    p_highlight = doc.add_paragraph('TYPO-HL-YELLOW-001 ')
    run_yellow = p_highlight.add_run('Yellow highlight')
    run_yellow.font.highlight_color = WD_COLOR_INDEX.YELLOW

    p_highlight.add_run(', TYPO-HL-GREEN-001 ')
    run_hl_green = p_highlight.add_run('Green highlight')
    run_hl_green.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN

    p_highlight.add_run(', TYPO-HL-CYAN-001 ')
    run_cyan = p_highlight.add_run('Cyan highlight')
    run_cyan.font.highlight_color = WD_COLOR_INDEX.TURQUOISE

    # Section 5: Font Limitations
    doc.add_heading('TYPO-LIMITS-001 Font Limitations', level=2)

    doc.add_paragraph(
        'TYPO-FONT-FAMILY-001 Font family names are NOT preserved by pandoc. '
        'This text has a specific font family (Arial) but will appear as default '
        'in markdown output.'
    )

    doc.add_paragraph(
        'TYPO-FONT-SIZE-001 Font sizes are also NOT preserved by pandoc. '
        'This limitation is documented in tests.'
    )

    # Section 6: Expected Behavior
    doc.add_heading('Expected Behavior', level=2)

    doc.add_paragraph(
        'When converting to Markdown, pandoc should:'
    )

    doc.add_paragraph(
        '1. Convert heading levels to # syntax'
    )

    doc.add_paragraph(
        '2. Convert text colors to HTML <span style="color:..."> tags'
    )

    doc.add_paragraph(
        '3. Convert alignment to HTML attributes or styles'
    )

    doc.add_paragraph(
        '4. Convert highlighting to HTML <span> with background colors'
    )

    doc.add_paragraph(
        '5. NOT preserve font family or size (known limitation)'
    )

    # Save the fixture
    output_path = Path(__file__).resolve().parents[1] / 'tests' / 'fixtures' / 'converter' / 'typography_sample.docx'
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f'✅ Created: {output_path}')


if __name__ == '__main__':
    main()
