#!/usr/bin/env python3
"""Create a DOCX fixture to test bookmark and anchor preservation.

Bookmarks are converted by pandoc as:
- Named bookmarks → HTML anchor tags (<a name="bookmark_name">)
- Cross-references → Links to anchors
- Bookmark text content preserved

This fixture creates a document with multiple bookmarks and links to them.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def add_bookmark(paragraph, bookmark_name: str, text: str):
    """Add a bookmark to a paragraph run.

    Bookmarks require low-level XML manipulation in python-docx.
    """
    run = paragraph.add_run(text)

    # Create bookmark start element
    bookmark_start = OxmlElement('w:bookmarkStart')
    bookmark_start.set(qn('w:id'), '0')
    bookmark_start.set(qn('w:name'), bookmark_name)

    # Create bookmark end element
    bookmark_end = OxmlElement('w:bookmarkEnd')
    bookmark_end.set(qn('w:id'), '0')

    # Insert bookmark around the run
    run._element.addprevious(bookmark_start)
    run._element.addnext(bookmark_end)

    return run


def main():
    """Create bookmarks_sample.docx fixture."""
    doc = Document()

    # Title
    doc.add_heading('BOOKMARK-TEST-001 Document Bookmarks Test', level=1)

    doc.add_paragraph(
        'This document tests bookmark preservation through DOCX→Markdown conversions.'
    )

    doc.add_paragraph(
        'Pandoc converts bookmarks to HTML anchor tags, which allows '
        'cross-references to work in the output.'
    )

    # Section 1: Named bookmarks
    doc.add_heading('BOOKMARK-NAMED-001 Named Bookmarks', level=2)

    p1 = doc.add_paragraph('This paragraph contains a bookmark: ')
    add_bookmark(p1, 'section_intro', 'BOOKMARK-ANCHOR-001 Introduction Section')

    doc.add_paragraph(
        'Named bookmarks allow you to create anchor points in the document '
        'that can be referenced from other locations.'
    )

    # Section 2: Multiple bookmarks
    doc.add_heading('BOOKMARK-MULTI-001 Multiple Bookmarks', level=2)

    p2 = doc.add_paragraph('First bookmark: ')
    add_bookmark(p2, 'bookmark_one', 'BOOKMARK-ONE-001 First Anchor')

    p3 = doc.add_paragraph('Second bookmark: ')
    add_bookmark(p3, 'bookmark_two', 'BOOKMARK-TWO-001 Second Anchor')

    p4 = doc.add_paragraph('Third bookmark: ')
    add_bookmark(p4, 'bookmark_three', 'BOOKMARK-THREE-001 Third Anchor')

    # Section 3: Cross-references (simulated with hyperlinks)
    doc.add_heading('BOOKMARK-XREF-001 Cross-References', level=2)

    doc.add_paragraph(
        'Links to bookmarks create cross-references within the document:'
    )

    # Add hyperlinks to the bookmarks (simulated)
    p5 = doc.add_paragraph('• See ')
    p5.add_run('section_intro').bold = True
    p5.add_run(' above')

    p6 = doc.add_paragraph('• See ')
    p6.add_run('bookmark_one').bold = True
    p6.add_run(', ')
    p6.add_run('bookmark_two').bold = True
    p6.add_run(', and ')
    p6.add_run('bookmark_three').bold = True

    # Section 4: Expected behavior
    doc.add_heading('Expected Behavior', level=2)

    doc.add_paragraph(
        'When converting to Markdown, pandoc should:'
    )

    doc.add_paragraph(
        '1. Convert bookmarks to HTML anchor tags (<a name="...">)'
    )

    doc.add_paragraph(
        '2. Preserve bookmark names'
    )

    doc.add_paragraph(
        '3. Maintain bookmark text content'
    )

    # Save the fixture
    output_path = Path(__file__).resolve().parents[1] / 'tests' / 'fixtures' / 'converter' / 'bookmarks_sample.docx'
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f'✅ Created: {output_path}')


if __name__ == '__main__':
    main()
