#!/usr/bin/env python3
"""Create DOCX fixture for text boxes and shapes testing.

Text boxes and shapes have LIMITED pandoc support:
- Text box content → May be extracted as plain text
- Shape positioning → Lost
- Shape styling (colors, borders) → Lost
- Grouped shapes → May be ungrouped

⚠️ Will be implemented via Google Cloud with LibreOffice/ImageMagick.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document


def main() -> None:
    """Create textboxes_shapes_sample.docx fixture.

    The semantic markers in this document are used by
    tests/test_convert_backend_textboxes_shapes.py to verify that
    text content from text boxes and shape-related sections is
    preserved in DOCX→Markdown conversion, even though layout and
    styling are lost.
    """
    doc = Document()

    doc.add_heading("TEXTBOX-TEST-001 Text Boxes & Shapes Test", level=1)

    doc.add_paragraph(
        "This document tests text box and shape preservation for future "
        "Google Cloud implementation."
    )

    # Text box section
    doc.add_heading("TEXTBOX-CONTENT-001 Text Box Content", level=2)
    doc.add_paragraph(
        "TEXTBOX-TEXT-001 Below is a text box: "
        "[Text boxes would be manually created in Word with Insert > Text Box]"
    )
    doc.add_paragraph("TEXTBOX-INSIDE-001 This text would be inside a text box.")

    # Shapes section
    doc.add_heading("SHAPE-BASIC-001 Basic Shapes", level=2)
    doc.add_paragraph(
        "SHAPE-TEXT-001 Basic shapes (rectangles, circles, arrows) would be "
        "inserted here. python-docx has limited shape support, so these "
        "would be manually created in Word."
    )

    # Positioned content
    doc.add_heading("POSITION-TEST-001 Positioned Elements", level=2)
    doc.add_paragraph(
        "POSITION-TEXT-001 Elements with absolute positioning would appear "
        "here. Positioning information is typically lost in conversion."
    )

    # Grouped shapes
    doc.add_heading("GROUP-TEST-001 Grouped Shapes", level=2)
    doc.add_paragraph(
        "GROUP-TEXT-001 Multiple shapes grouped together would appear here. "
        "Grouping is often lost; shapes become individual elements."
    )

    # Text box with formatting
    doc.add_heading("TEXTBOX-FORMAT-001 Formatted Text Box", level=2)
    p = doc.add_paragraph("TEXTBOX-FORMAT-TEXT-001 ")
    p.add_run("This text box has bold").bold = True
    p.add_run(" and ")
    p.add_run("italic").italic = True
    p.add_run(" formatting inside it.")

    # Expected behavior
    doc.add_heading("Expected Behavior", level=2)

    doc.add_paragraph(
        "⚠️ PANDOC LIMITATION: Text box content may be extracted, but "
        "positioning and shape styling are lost."
    )

    doc.add_paragraph(
        "✅ FUTURE: Will be implemented via Google Cloud with LibreOffice "
        "HTML export to preserve positioning and convert shapes to "
        "SVG/images."
    )

    doc.add_paragraph(
        "Tests validate: 1) Text content is extracted, 2) Markers present, "
        "3) Infrastructure ready for GCloud shape/positioning "
        "implementation."
    )

    output_path = (
        Path(__file__).resolve().parents[1]
        / "tests"
        / "fixtures"
        / "converter"
        / "textboxes_shapes_sample.docx"
    )
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"✅ Created: {output_path}")
    print("⚠️  NOTE: For full testing, manually add text boxes and shapes in Word.")


if __name__ == "__main__":
    main()

