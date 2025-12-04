#!/usr/bin/env python3
"""Create DOCX fixture for paragraph formatting (indents, spacing, line spacing).

These features have LIMITED pandoc support but will be implemented via Google Cloud Document AI.
This fixture establishes test infrastructure for future GCloud implementation.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


def main():
    """Create paragraph_formatting_sample.docx fixture."""
    doc = Document()

    doc.add_heading('PARA-FORMAT-TEST-001 Paragraph Formatting Test', level=1)

    doc.add_paragraph(
        'This document tests paragraph formatting features for future Google Cloud implementation.'
    )

    # First-line indent
    doc.add_heading('PARA-INDENT-FIRST-001 First-Line Indent', level=2)
    p1 = doc.add_paragraph('INDENT-FIRST-TEXT-001 This paragraph has a first-line indent. '
                           'The first line should be indented 0.5 inches from the left margin.')
    p1.paragraph_format.first_line_indent = Inches(0.5)

    # Hanging indent
    doc.add_heading('PARA-INDENT-HANG-001 Hanging Indent', level=2)
    p2 = doc.add_paragraph('INDENT-HANG-TEXT-001 This paragraph has a hanging indent. '
                           'The first line starts at the margin, but subsequent lines are indented.')
    p2.paragraph_format.left_indent = Inches(0.5)
    p2.paragraph_format.first_line_indent = Inches(-0.5)

    # Left/Right margins
    doc.add_heading('PARA-MARGINS-001 Left and Right Margins', level=2)
    p3 = doc.add_paragraph('MARGINS-TEXT-001 This paragraph has both left and right indentation, '
                           'creating narrower text margins.')
    p3.paragraph_format.left_indent = Inches(1.0)
    p3.paragraph_format.right_indent = Inches(1.0)

    # Paragraph spacing (before/after)
    doc.add_heading('PARA-SPACING-001 Paragraph Spacing', level=2)
    p4 = doc.add_paragraph('SPACING-BEFORE-001 This paragraph has space before.')
    p4.paragraph_format.space_before = Pt(24)

    p5 = doc.add_paragraph('SPACING-AFTER-001 This paragraph has space after.')
    p5.paragraph_format.space_after = Pt(24)

    p6 = doc.add_paragraph('SPACING-BOTH-001 This paragraph has space before AND after.')
    p6.paragraph_format.space_before = Pt(12)
    p6.paragraph_format.space_after = Pt(12)

    # Line spacing
    doc.add_heading('PARA-LINE-SPACING-001 Line Spacing', level=2)

    p7 = doc.add_paragraph('LINE-SINGLE-001 This paragraph has single line spacing (1.0). '
                           'This is the default spacing.')
    p7.paragraph_format.line_spacing = 1.0

    p8 = doc.add_paragraph('LINE-ONEHALF-001 This paragraph has 1.5 line spacing. '
                           'There is more space between lines.')
    p8.paragraph_format.line_spacing = 1.5

    p9 = doc.add_paragraph('LINE-DOUBLE-001 This paragraph has double line spacing (2.0). '
                           'There is even more space between lines.')
    p9.paragraph_format.line_spacing = 2.0

    # Expected behavior
    doc.add_heading('Expected Behavior', level=2)

    doc.add_paragraph(
        '⚠️ PANDOC LIMITATION: These formatting features are NOT preserved by pandoc.'
    )

    doc.add_paragraph(
        '✅ FUTURE: Will be implemented via Google Cloud Document AI extraction.'
    )

    doc.add_paragraph(
        'Tests validate: 1) Content is preserved, 2) Markers are present, '
        '3) Test infrastructure ready for GCloud implementation.'
    )

    output_path = Path(__file__).resolve().parents[1] / 'tests' / 'fixtures' / 'converter' / 'paragraph_formatting_sample.docx'
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f'✅ Created: {output_path}')


if __name__ == '__main__':
    main()
