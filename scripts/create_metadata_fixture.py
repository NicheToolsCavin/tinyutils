#!/usr/bin/env python3
"""Create DOCX fixture for document metadata testing.

Document metadata includes:
- Core properties: Title, Subject, Author, Keywords, Category, Comments
- Extended properties: Company, Manager, etc.
- Document statistics: Word count, page count, etc.

Pandoc behavior:
- Core properties → May appear in YAML front matter (with --standalone)
- Custom properties → Not extracted by default
- Statistics → Not preserved

This fixture creates a document with core metadata properties.
"""
from __future__ import annotations

from pathlib import Path
from datetime import datetime

from docx import Document


def main():
    """Create metadata_sample.docx fixture."""
    doc = Document()

    # Set core document properties
    core_props = doc.core_properties
    core_props.title = "TinyUtils Test Document"
    core_props.subject = "Document Conversion Testing"
    core_props.author = "Claude Code Agent"
    core_props.keywords = "testing; conversion; pandoc; tinyutils"
    core_props.comments = "This is a test document for metadata preservation testing."
    core_props.category = "Testing"
    core_props.created = datetime.now()
    core_props.modified = datetime.now()

    # Add document content explaining metadata
    doc.add_heading('METADATA-TEST-001 Document Metadata Test', level=1)

    doc.add_paragraph(
        'This document tests metadata preservation through DOCX conversions.'
    )

    # Core properties
    doc.add_heading('Core Properties', level=2)

    doc.add_paragraph('META-TITLE-001 Title: TinyUtils Test Document')
    doc.add_paragraph('META-SUBJECT-001 Subject: Document Conversion Testing')
    doc.add_paragraph('META-AUTHOR-001 Author: Claude Code Agent')
    doc.add_paragraph('META-KEYWORDS-001 Keywords: testing; conversion; pandoc; tinyutils')
    doc.add_paragraph('META-CATEGORY-001 Category: Testing')
    doc.add_paragraph('META-COMMENTS-001 Comments: This is a test document for metadata preservation testing.')

    # Expected behavior
    doc.add_heading('Expected Behavior', level=2)

    doc.add_paragraph(
        'Pandoc behavior: Core properties may be extracted to YAML front matter '
        'when using --standalone flag.'
    )

    doc.add_paragraph(
        'Note: These metadata markers in the document body are ALWAYS preserved, '
        'regardless of whether pandoc extracts the document properties.'
    )

    doc.add_paragraph(
        'Tests validate: 1) Conversion succeeds with metadata, '
        '2) Document body content is preserved, '
        '3) Metadata markers are present in output.'
    )

    output_path = Path(__file__).resolve().parents[1] / 'tests' / 'fixtures' / 'converter' / 'metadata_sample.docx'
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f'✅ Created: {output_path}')


if __name__ == '__main__':
    main()
