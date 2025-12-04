#!/usr/bin/env python3
"""Create a DOCX fixture to test comments and annotations.

Comments are converted by pandoc with LIMITED support:
- Comments can be extracted with special options
- By default, comments are NOT preserved
- Comment text content may be lost
- Document text IS preserved

This fixture creates a document with comments for testing.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn


def add_comment_to_run(run, author: str, text: str, comment_id: int):
    """Add a comment to a run (requires low-level XML manipulation)."""
    # Note: python-docx doesn't have built-in comment support
    # We'll create markers that indicate where comments should be
    # The actual comment structure is complex in OOXML
    run.add_text(f" [COMMENT-{comment_id:03d} by {author}: {text}]")


def main():
    """Create comments_sample.docx fixture."""
    doc = Document()

    # Title
    doc.add_heading('COMMENTS-TEST-001 Document Comments Test', level=1)

    doc.add_paragraph(
        'This document tests comment preservation through DOCX→Markdown conversions.'
    )

    doc.add_paragraph(
        'Note: python-docx has limited comment support, so comments are simulated '
        'with inline markers.'
    )

    # Section 1: Inline Comments
    doc.add_heading('COMMENTS-INLINE-001 Inline Comments', level=2)

    p1 = doc.add_paragraph('COMMENT-TEXT-001 This text has a comment')
    add_comment_to_run(p1.runs[0], "Alice", "This needs clarification", 1)

    doc.add_paragraph(
        'Comments are typically used for document review and collaboration.'
    )

    # Section 2: Multiple Comments
    doc.add_heading('COMMENTS-MULTI-001 Multiple Comments', level=2)

    p2 = doc.add_paragraph('COMMENT-MULTI-A-001 First commented text')
    add_comment_to_run(p2.runs[0], "Bob", "Good point", 2)

    p3 = doc.add_paragraph('COMMENT-MULTI-B-001 Second commented text')
    add_comment_to_run(p3.runs[0], "Carol", "Needs revision", 3)

    p4 = doc.add_paragraph('COMMENT-MULTI-C-001 Third commented text')
    add_comment_to_run(p4.runs[0], "Dave", "Approved", 4)

    # Section 3: Comment Threads
    doc.add_heading('COMMENTS-THREAD-001 Comment Threads', level=2)

    p5 = doc.add_paragraph('COMMENT-THREAD-START-001 This starts a discussion')
    add_comment_to_run(p5.runs[0], "Alice", "I disagree with this", 5)
    p5.add_run(' ')
    add_comment_to_run(p5.runs[-1], "Bob", "Reply: Why do you disagree?", 6)

    # Section 4: Expected Behavior
    doc.add_heading('Expected Behavior', level=2)

    doc.add_paragraph(
        'When converting to Markdown, pandoc typically:'
    )

    doc.add_paragraph(
        '1. Does NOT preserve comments by default'
    )

    doc.add_paragraph(
        '2. Preserves the underlying text content'
    )

    doc.add_paragraph(
        '3. May extract comments with special filter options'
    )

    doc.add_paragraph(
        '4. Comment markers in this fixture will be visible in output'
    )

    # Save the fixture
    output_path = Path(__file__).resolve().parents[1] / 'tests' / 'fixtures' / 'converter' / 'comments_sample.docx'
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f'✅ Created: {output_path}')


if __name__ == '__main__':
    main()
