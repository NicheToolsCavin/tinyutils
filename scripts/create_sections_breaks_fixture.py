#!/usr/bin/env python3
"""Create a DOCX fixture to test section breaks and page breaks.

Sections and breaks are converted by pandoc as:
- Page breaks → Preserved as horizontal rules or raw blocks
- Section breaks → May preserve some properties
- Different page orientations per section → Limited preservation

This fixture creates a document with multiple sections and break types.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.shared import Pt


def main():
    """Create sections_breaks_sample.docx fixture."""
    doc = Document()

    # Title
    doc.add_heading('SECTIONS-TEST-001 Document Sections and Breaks Test', level=1)

    doc.add_paragraph(
        'This document tests section breaks and page breaks through DOCX→Markdown conversions.'
    )

    # Section 1: Page Breaks
    doc.add_heading('BREAKS-PAGE-001 Page Breaks', level=2)

    doc.add_paragraph(
        'BREAK-BEFORE-001 This paragraph appears before the page break.'
    )

    # Add page break
    doc.add_page_break()

    doc.add_paragraph(
        'BREAK-AFTER-001 This paragraph appears after the page break. '
        'Page breaks should be preserved as horizontal rules or similar markers.'
    )

    # Section 2: Continuous Section Break
    doc.add_heading('SECTIONS-CONTINUOUS-001 Continuous Section Break', level=2)

    doc.add_paragraph(
        'SECTION-CONT-BEFORE-001 This text is before a continuous section break. '
        'Continuous breaks allow format changes without starting a new page.'
    )

    # Add continuous section break
    section = doc.add_section(WD_SECTION.CONTINUOUS)
    section.start_type = WD_SECTION.CONTINUOUS

    doc.add_paragraph(
        'SECTION-CONT-AFTER-001 This text is after the continuous section break. '
        'It should be on the same page but in a new section.'
    )

    # Section 3: Next Page Section Break
    doc.add_heading('SECTIONS-NEXT-PAGE-001 Next Page Section Break', level=2)

    doc.add_paragraph(
        'SECTION-NP-BEFORE-001 This text is before a next-page section break.'
    )

    # Add next page section break
    section = doc.add_section(WD_SECTION.NEW_PAGE)

    doc.add_paragraph(
        'SECTION-NP-AFTER-001 This text appears after the next-page section break. '
        'It should start on a new page.'
    )

    # Section 4: Even/Odd Page Breaks
    doc.add_heading('SECTIONS-EVEN-ODD-001 Even and Odd Page Breaks', level=2)

    doc.add_paragraph(
        'SECTION-EVEN-001 Testing even page section break.'
    )

    # Add even page section break
    section = doc.add_section(WD_SECTION.EVEN_PAGE)

    doc.add_paragraph(
        'SECTION-EVEN-AFTER-001 This should start on an even-numbered page.'
    )

    doc.add_paragraph(
        'SECTION-ODD-001 Testing odd page section break.'
    )

    # Add odd page section break
    section = doc.add_section(WD_SECTION.ODD_PAGE)

    doc.add_paragraph(
        'SECTION-ODD-AFTER-001 This should start on an odd-numbered page.'
    )

    # Section 5: Multiple Page Breaks
    doc.add_heading('BREAKS-MULTI-001 Multiple Page Breaks', level=2)

    doc.add_paragraph(
        'BREAK-MULTI-1-001 First paragraph.'
    )

    doc.add_page_break()

    doc.add_paragraph(
        'BREAK-MULTI-2-001 Second paragraph after first break.'
    )

    doc.add_page_break()

    doc.add_paragraph(
        'BREAK-MULTI-3-001 Third paragraph after second break.'
    )

    # Section 6: Expected Behavior
    doc.add_heading('Expected Behavior', level=2)

    doc.add_paragraph(
        'When converting to Markdown, pandoc should:'
    )

    doc.add_paragraph(
        '1. Preserve page breaks as horizontal rules (---) or similar markers'
    )

    doc.add_paragraph(
        '2. Preserve section break markers in some form'
    )

    doc.add_paragraph(
        '3. Maintain content sequence across breaks'
    )

    doc.add_paragraph(
        '4. Keep all text content even if formatting/pagination is lost'
    )

    # Save the fixture
    output_path = Path(__file__).resolve().parents[1] / 'tests' / 'fixtures' / 'converter' / 'sections_breaks_sample.docx'
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f'✅ Created: {output_path}')


if __name__ == '__main__':
    main()
