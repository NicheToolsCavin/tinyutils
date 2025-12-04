#!/usr/bin/env python3
"""Create DOCX fixture for cross-reference testing.

Cross-references have LIMITED pandoc support:
- REF fields → Static text only
- PAGEREF fields → Static page numbers only
- Cross-reference functionality → Lost

⚠️ Will be implemented via Google Cloud with LibreOffice field extraction.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def add_cross_reference(paragraph, ref_id: str, text: str, ref_type: str = "heading"):
    """Add a cross-reference field to a paragraph.

    Args:
        paragraph: The paragraph to add the reference to
        ref_id: The bookmark ID being referenced
        text: The display text
        ref_type: Type of reference (heading, figure, table)
    """
    run = paragraph.add_run(text)

    # Create REF field
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = f' REF {ref_id} \\h '

    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')

    run._element.addprevious(fldChar_begin)
    run._element.addprevious(instrText)
    run._element.addnext(fldChar_end)

    return run


def main():
    """Create cross_references_sample.docx fixture."""
    doc = Document()

    doc.add_heading('XREF-TEST-001 Cross-References Test', level=1)

    doc.add_paragraph(
        'This document tests cross-reference preservation for future Google Cloud implementation.'
    )

    # Section 1: Reference to heading
    doc.add_heading('XREF-HEADING-001 Section One', level=2)
    p1 = doc.add_paragraph('XREF-HEADING-TEXT-001 This is section one content.')

    # Section 2: Reference to figure
    doc.add_heading('XREF-FIGURE-001 Figure Reference', level=2)
    doc.add_paragraph('XREF-FIGURE-TEXT-001 See Figure 1 below:')
    doc.add_paragraph('XREF-FIGURE-CAPTION-001 Figure 1: Sample Figure')

    # Section 3: Reference to table
    doc.add_heading('XREF-TABLE-001 Table Reference', level=2)
    doc.add_paragraph('XREF-TABLE-TEXT-001 See Table 1 below:')
    doc.add_paragraph('XREF-TABLE-CAPTION-001 Table 1: Sample Table')

    # Section 4: Cross-references to above items
    doc.add_heading('XREF-LINKS-001 Cross-Reference Links', level=2)

    p_ref1 = doc.add_paragraph('XREF-LINK-TEXT-001 See section: ')
    add_cross_reference(p_ref1, 'xref_heading_001', 'Section One', 'heading')

    p_ref2 = doc.add_paragraph('XREF-LINK-TEXT-002 See figure: ')
    add_cross_reference(p_ref2, 'xref_figure_001', 'Figure 1', 'figure')

    p_ref3 = doc.add_paragraph('XREF-LINK-TEXT-003 See table: ')
    add_cross_reference(p_ref3, 'xref_table_001', 'Table 1', 'table')

    # PAGEREF example
    doc.add_heading('XREF-PAGEREF-001 Page References', level=2)
    p_pageref = doc.add_paragraph('XREF-PAGEREF-TEXT-001 See page: ')
    p_pageref.add_run('Page 1').bold = True
    p_pageref.add_run(' (This would be a PAGEREF field in Word)')

    # Expected behavior
    doc.add_heading('Expected Behavior', level=2)

    doc.add_paragraph(
        '⚠️ PANDOC LIMITATION: Cross-reference fields → Static text only. '
        'REF/PAGEREF field codes and linking functionality are lost.'
    )

    doc.add_paragraph(
        '✅ FUTURE: Will be implemented via Google Cloud with LibreOffice '
        'to extract REF fields and create proper HTML links.'
    )

    output_path = Path(__file__).resolve().parents[1] / 'tests' / 'fixtures' / 'converter' / 'cross_references_sample.docx'
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f'✅ Created: {output_path}')


if __name__ == '__main__':
    main()
